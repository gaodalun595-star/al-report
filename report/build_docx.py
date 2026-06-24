#!/usr/bin/env python
"""手動組裝專題報告.docx：嵌入 server、architecture、demo 全部 PNG。"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPORT = Path("report")
ASSETS = REPORT / "assets"
MD = REPORT / "專題報告.md"
OUT = REPORT / "專題報告.docx"

IMG_PATTERN = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
HEADING_PATTERN = re.compile(r"^(#{1,6})\s+(.+)$")
BULLET_PATTERN = re.compile(r"^[-*]\s+(.+)$")
META_PATTERN = re.compile(r"^\*\*(.+?)\*\*$")
SEP = "---"

doc = Document()

def add_image(rel_path: str):
    abs_path = (REPORT / rel_path).resolve()
    if not abs_path.exists():
        return False
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(abs_path), width=Cm(15))
    return True

with MD.open(encoding="utf-8") as f:
    lines = f.read().splitlines()

in_code = False
for line in lines:
    if line.strip().startswith("```"):
        in_code = not in_code
        continue
    if in_code:
        continue

    if line.strip() == SEP:
        doc.add_paragraph()
        continue

    m = HEADING_PATTERN.match(line)
    if m:
        level = len(m.group(1))
        text = m.group(2).strip()
        doc.add_heading(text, level=min(level, 4))
        continue

    m = IMG_PATTERN.search(line)
    if m:
        add_image(m.group(2).strip())
        continue

    m = META_PATTERN.match(line.strip())
    if m:
        doc.add_paragraph(m.group(1).strip())
        continue

    m = BULLET_PATTERN.match(line)
    if m:
        doc.add_paragraph(m.group(1).strip(), style="List Bullet")
        continue

    if line.strip():
        doc.add_paragraph(line.strip())

doc.save(OUT)
print(f"OK: {OUT}")
